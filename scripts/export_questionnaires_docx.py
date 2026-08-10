"""Export the questionnaire specification to a mobile-friendly Word document."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "QUESTIONNAIRES.md"
OUTPUT = ROOT / "docs" / "Recon实验问卷方案.docx"


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(36, 95, 159)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def is_separator(row: str) -> bool:
    cells = [item.strip() for item in row.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", item) for item in cells)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (
        ("Title", 22, "183B59"),
        ("Heading 1", 17, "183B59"),
        ("Heading 2", 14, "245F9F"),
        ("Heading 3", 12, "245F9F"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = styles["Normal"]
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                paragraph.paragraph_format.left_indent = Cm(0.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                if not is_separator(lines[index]):
                    rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    add_inline(paragraph, value)
                    if row_index == 0:
                        set_cell_shading(cell, "DCE9F5")
                        for run in paragraph.runs:
                            run.bold = True
            document.add_paragraph()
            continue
        if not stripped or stripped == "---":
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and index == 0:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, text)
                subtitle = document.add_paragraph("Trial、Condition 与最终比较问卷")
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle.runs[0].font.color.rgb = RGBColor(101, 117, 134)
            else:
                add_inline(document.add_heading(level=min(level, 3)), text)
            index += 1
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            run = paragraph.add_run(stripped.lstrip("> "))
            run.italic = True
            run.font.color.rgb = RGBColor(70, 86, 101)
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            add_inline(document.add_paragraph(style="List Bullet"), bullet.group(1))
        elif numbered:
            add_inline(document.add_paragraph(style="List Number"), numbered.group(1))
        else:
            add_inline(document.add_paragraph(), stripped)
        index += 1

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Recon EEG + Gaze 实验问卷方案").font.size = Pt(8)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
